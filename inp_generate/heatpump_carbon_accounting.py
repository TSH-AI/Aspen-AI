import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import asyncio

os.environ["SERPER_API_KEY"] = "<your api key>"

# Add rag_modular directory to Python path
sys.path.append(os.path.join(os.path.dirname(os.path.dirname(__file__)), 'rag_modular'))

# Add rag_table directory to Python path
sys.path.append(os.path.join(os.path.dirname(os.path.dirname(__file__)), 'rag_table'))

from langchain_experimental.tools import PythonREPLTool
from langchain.agents import initialize_agent, Tool
from langchain_community.utilities import GoogleSerperAPIWrapper
from langchain_openai import ChatOpenAI
from power_range_rag_local import get_power_range_from_rag
from gxt_rag import get_gxt_table

llm = ChatOpenAI(model="<your model>", api_key="<your api key>",
                   base_url="<your base url>", temperature=0)

search_tool = GoogleSerperAPIWrapper()
python_repl_tool = PythonREPLTool()

tools = [
    Tool(
        name="GoogleSerper",
        func=search_tool.run,
        description="Use GoogleSerper to search for real-time internet information."
    ),
    Tool(
        name="Python_REPL",
        func=python_repl_tool.run,
        description="Execute Python code for mathematical or programming tasks."
    )
]

# Format instructions for agent initialization
FORMAT_INSTRUCTIONS = """You MUST follow this format:

Thought: Steps to solve the problem
Action: Tool to use (must be one of [GoogleSerper, Python_REPL])
Action Input: Tool input content
Observation: Tool return result
... (this loop can repeat multiple times)
Final Answer: Final answer (use when task is complete)"""

agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent="zero-shot-react-description",
    verbose=True,
    handle_parsing_errors=True,
    return_intermediate_steps=True,
    agent_kwargs={
        'format_instructions': FORMAT_INSTRUCTIONS
    }
)

def save_to_word(content: str, filename: str = None):
    """Save the content to a Word document"""
    # Dynamically generate output directory path based on current date
    current_date = datetime.now().strftime("%b %d").lower()
    base_dir = "inp_generate/difficult/carbon_accounting"
    output_dir = base_dir
    
    if filename is None:
        filename = "heatpump_carbon_emissions.docx"
    
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Carbon Emissions Analysis Report for Heat Pump Design', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()
    
    # Add content
    doc.add_paragraph(content)
    
    # Save the document
    doc.save(filepath)
    print(f"\n\033[1;32mReport saved as: {filepath}\033[0m")
    return filepath

async def setup_factors():
    """Setup emission factors"""
    global power_factors
    power_factors = await get_power_range_from_rag()

gxt_table_data = asyncio.run(get_gxt_table())

async def main():
    # Initialize emission factors
    await setup_factors()
    
    # Dynamically generate REP file path based on current date
    current_date = datetime.now().strftime("%b %d").lower()
    base_dir = "inp_generate/difficult"
    rep_file = os.path.join(base_dir, f"{current_date}heat_pump_design.rep")
    
    print(f"Using heat pump design report file: {rep_file}")
    
    # Check if file exists
    if not os.path.exists(rep_file):
        raise FileNotFoundError(f"Heat pump design report file not found: {rep_file}")
    
    with open(rep_file, "r", encoding='utf-8') as f:
        rep_content = f.read()

    # Single prompt for carbon emissions calculation
    prompt = f"""
You are a senior chemical process engineer. Utilize maximum model capacity to execute the following tasks:

Here are some general instructions for cooling water flowrate calculation. If updated information is not given, follow the default settings.
  - Default temperature rise is 5°C
  - According to the condenser duty in the .rep file

Here are some general instructions for the distillation column's condenser height calculation. If updated information is not given, follow the default settings.
  - Assume that the height of each plate of the distillation column is 0.5 m, multiply the number of plates to get the total height of the distillation column.
  - Remember that the bottom of the distillation column is 1 m above the ground.
  - The condenser is 3 m above the distillation tower top.

Here are some general instructions for pump work calculation. If updated information is not given, follow the default settings.
  - The cooling water tower is set on the ground
  - Three steps to calculate the diameter and fluid velocity. Display the calculation process. Assume the pump inlet pipe has the same diameter as the outlet pipe, which is calculated below.
    1. When calculating the pipe diameter, assume 2 m/s as the initial velocity for liquid and 20m/s for gas. 
    2. Next round the diameter up to a standard pipe diameter: inner diameters: 15 mm, 20 mm, 25 mm, 32 mm, 40 mm, 45 mm, 50 mm, 70 mm, 81 mm, 100 mm, 125 mm, 149 mm）. 
    3. Recalculate the velocity based on the new diameter. Doublecheck that the final velocity should be 1-3 m/s for liquid, 10-30m/s for gas
  - The diameter and fluid velocity should be calculated before the head loss calculation because the value of u1 and u2 is needed.
  - Pay attention to the pressure difference, especially for the plus or minus sign.
  - Consider all energy losses in the system using Bernoulli's equation.
  - Consider the flow type: turbulent or laminar, before calculating the friction factor. Use Blasius equation for turbulent flow and Poiseuille equation for laminar flow.
  - Use the following local resistance coefficient: 90 degree elbow: 0.75; 45 degree elbow: 0.35; 180 degree elbow: 1.5; tee: 1; search RAG for other local resistance coefficient.

Here are some general instruction for steam consumption calculation. If updated information is not given, follow the default settings.
  - Default latent heat is 2200kJ/kg. 

Here are some general instructions for carbon emissions calculation. If updated information is not given, follow the default settings.
  - The steam consumption is related to the heat duty of the reboiler.

Some specific requirements for this pump work calculation:      
  - Pressure at the cooling water pump inlet: 1 atm
  - Pressure at the condenser inlet pipe: 10 kPa, because the cooling water falls down and the flow is not continuous any more (Industrial experience). 
  - Straight pipe resistance (40m length)
  - 5 90 degree elbows
  - (Search RAG for a typical pump efficiency) Use 0.6 as the default efficiency.
  - The cooling water tower is set on the ground

Calculate carbon emissions for heat pump design and provide the complete detailed calculation process in the finished chain of thought, including formula, calculation and result:

IMPORTANT: You MUST show ALL calculation steps, formulas, and intermediate results. Do not skip any calculations. Use the Python_REPL tool for all mathematical operations and unit conversions. Double-check all calculations for accuracy, especially unit conversions and numerical computations.

Any step that requires calculation should be done using the Python_REPL tool, including the unit conversion.
If any data is not available, please search online for the data and cite the source, especially for the energy consumption of the cooling tower's fan.

The power-consuming device of the whole system are the compressor, the cooling tower's fan and the cooling water pump for heater2:
1. Utilize the Bernoulli's equation to calculate the energy consumption from the pump to the cooling tower for heater2.
1.1 Calculate the cooling water flow rate (5°C ΔT)
1.2 Calculate total pump head requirements using resistance loss methods and the Bernoulli's equation:
   - Straight pipe resistance (40m length):
     * Calculate the pipe diameter according to the cooling water flow rate.
     * Next round the diameter up to a standard pipe diameter: inner diameters: 15 mm, 20 mm, 25 mm, 32 mm, 40 mm, 45 mm, 50 mm, 70 mm, 81 mm, 100 mm, 125 mm, 149 mm）. 
     * Recalculate the flow velocity based on the new pipe diameter before proceeding with subsequent calculations.
     * Calculate Reynolds number for flow regime
     * Use Blasius equation with friction factor according to the Reynolds number
   - Local resistance:
     * 5 elbows (use resistance coefficient=0.75)
     * Calculate each elbow's head loss
   - Liquid level resistance:
     * The heater2 is set on the ground
     * Search the height of the cooling tower from {gxt_table_data}
2. Calculate pump power using:
   - Total head (sum of all resistance losses)
   - Cooling water mass flow rate
   - Pump efficiency:0.6
3. The energy consumption of the cooling tower's fan is decided by the cooling water flow rate. Search following table for the energy consumption of the cooling tower's fan:{gxt_table_data} and convert the energy consumption at the ratio using current cooling water flow rate.
4. Extract the energy consumption of the compressor from the input data.
5. Power emission factors:{power_factors}
6. Calculate the carbon emissions for the compressor, the cooling tower's fan and the cooling water pump for heater2.
7. For each data point, provide:
   - Source of the value (if from input data)
   - Complete calculation process (if calculated)
   - All intermediate steps and formulas used
   - Units for all values and pay attention to the unit conversion
   - Verify all calculations using Python_REPL tool to ensure accuracy
8. Output format:
First provide the complete detailed calculation process including:
- All formulas used
- Step-by-step calculations with intermediate results
- Unit conversions
- All intermediate steps and reasoning

Then provide the final results table:
| Fuel Type | Power Emissions (kg/h) | Error Range |
|---------|-----------------|-----------------|
| Coal    |-----------------|-----------------|
| Natural Gas |-----------------|-----------------|
| Wind Power |-----------------|-----------------|

IMPORTANT: All calculations must be verified using Python_REPL tool. Pay special attention to:
- Unit conversions (especially mm to m conversions)
- Numerical accuracy in all mathematical operations
- Correct application of formulas and constants

Raw data:
{rep_content}

Note: {FORMAT_INSTRUCTIONS}"""

    print("\n\033[1;34m=== Executing Carbon Emissions Calculation ===\033[0m")
    
    try:
        response = agent.invoke({"input": prompt})
        raw_output = response['output']
        
        # Parse output
        parsed_output = raw_output
        
        print(f"\n\033[1;32mCalculation Result:\n{parsed_output}\033[0m")
        
        # Save to Word document
        save_to_word(parsed_output)
        
    except Exception as e:
        print(f"\033[1;31mCalculation failed: {str(e)}\033[0m")

if __name__ == "__main__":
    asyncio.run(main())