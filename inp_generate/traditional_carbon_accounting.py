import os
import re
from datetime import datetime
import webbrowser
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import asyncio

# Add rag_modular directory to Python path
sys.path.append(os.path.join(os.path.dirname(os.path.dirname(__file__)), 'rag_modular'))

# Add rag_table directory to Python path
sys.path.append(os.path.join(os.path.dirname(os.path.dirname(__file__)), 'rag_table'))

os.environ["SERPER_API_KEY"] = "<your api key>"

from langchain_experimental.agents.agent_toolkits import create_python_agent
from langchain_experimental.tools import PythonREPLTool
from langchain.agents import initialize_agent, Tool
from langchain_community.utilities import GoogleSerperAPIWrapper
from langchain_openai import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from steam_factors_rag_local import run_query_steam_factors
from power_value_rag_local import get_power_averages_from_rag
from gxt_rag import get_gxt_table

llm = ChatOpenAI(model="<your model>", api_key="<your api key>",
                 base_url="<your base url>", temperature=0)

search_tool = GoogleSerperAPIWrapper()
python_repl_tool = PythonREPLTool()

tools = [
    Tool(
        name="GoogleSerper",
        func=search_tool.run,
        description="Use GoogleSerper to search for real-time internet information."
    ),
    Tool(
        name="Python_REPL",
        func=python_repl_tool.run,
        description="Execute Python code for mathematical or programming tasks."
    )
]

# Format instructions for agent initialization
FORMAT_INSTRUCTIONS = """You MUST follow this format:

Thought: Steps to solve the problem
Action: Tool to use (must be one of [GoogleSerper, Python_REPL])
Action Input: Tool input content
Observation: Tool return result
... (this loop can repeat multiple times)
Final Answer: Final answer (use when task is complete)"""

agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent="zero-shot-react-description",
    verbose=True,
    handle_parsing_errors=True,
    return_intermediate_steps=True,
    agent_kwargs={
        'format_instructions': FORMAT_INSTRUCTIONS
    }
)

def save_to_word(content: str, filename: str = None):
    """Save the content to a Word document"""
    # Set the output directory
    output_dir = "inp_generate/simple/carbon_accounting"
    
    if filename is None:
        filename = "traditional_carbon_emissions.docx"
    
    # Create directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Combine directory and filename
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Carbon Emissions Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()
    
    # Add content
    doc.add_paragraph(content)
    
    # Save the document
    doc.save(filepath)
    print(f"\n\033[1;32mReport saved as: {filepath}\033[0m")
    return filepath

def try_read_file(file_path):
    """Try different encodings to read a file"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File does not exist: {file_path}\nPlease ensure the file path is correct, or place the file in the correct location.")
    
    encodings = ['utf-8', 'latin1', 'cp1252', 'gbk', 'gb18030', 'utf-16']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Cannot read file with the following encodings: {encodings}")

async def setup_factors():
    """Setup emission factors"""
    global steam_factors, power_factors
    print("\n\033[1;36m=== Retrieving RAG Data ===\033[0m")
    
    print("Retrieving steam emission factors...")
    steam_factors = await run_query_steam_factors()
    print(f"Steam emission factors retrieved: {steam_factors}")
    
    print("Retrieving power emission factors...")
    power_factors = await get_power_averages_from_rag()
    print(f"Power emission factors retrieved: {power_factors}")
    
    print("Retrieving GXT table data...")
    gxt_table_data = await get_gxt_table()
    print(f"GXT table data retrieved: {gxt_table_data}")
    
    print("\n\033[1;36m=== RAG Data Retrieval Complete ===\033[0m")
    return steam_factors, power_factors, gxt_table_data

async def main():
    # Set file paths
    current_dir = os.path.dirname(os.path.abspath(__file__))
    current_date = datetime.now().strftime("%b %d").lower()
    rep_file = os.path.join("inp_generate/simple", f"{current_date}radsimpleo.rep")
    inp_file = os.path.join("inp_generate/simple", f"{current_date}radsimpleo.inp")

    try:
        # Initialize emission factors
        steam_factors, power_factors, gxt_table_data = await setup_factors()
        
        try:
            rep_content = try_read_file(rep_file)
        except FileNotFoundError as e:
            print(f"\033[1;31mError: {str(e)}\033[0m")
            return
        
        try:
            inp_content = try_read_file(inp_file)
        except FileNotFoundError as e:
            print(f"\033[1;31mError: {str(e)}\033[0m")
            return

        # Single prompt for carbon emissions calculation
        prompt = f"""
You are a senior chemical process engineer. Utilize maximum model capacity to execute the following tasks:

Here are some general instructions for cooling water flowrate calculation. If updated information is not given, follow the default settings.
  - Default temperature rise is 5°C
  - According to the condenser duty in the .rep file

Here are some general instructions for the distillation column's condenser height calculation. If updated information is not given, follow the default settings.
  - Assume that the height of each plate of the distillation column is 0.5 m, multiply the number of plates to get the total height of the distillation column.
  - Remember that the bottom of the distillation column is 1 m above the ground.
  - The condenser is 3 m above the distillation tower top.

Here are some general instructions for pump work calculation. If updated information is not given, follow the default settings.
  - The cooling water tower is set on the ground
  - Three steps to calculate the diameter and fluid velocity. Display the calculation process. Assume the pump inlet pipe has the same diameter as the outlet pipe, which is calculated below.
    1. When calculating the pipe diameter, assume 2 m/s as the initial velocity for liquid and 20m/s for gas. 
    2. Next round the diameter up to a standard pipe diameter: inner diameters: 15 mm, 20 mm, 25 mm, 32 mm, 40 mm, 45 mm, 50 mm, 70 mm, 81 mm, 100 mm, 125 mm, 149 mm）. 
    3. Recalculate the velocity based on the new diameter. Doublecheck that the final velocity should be 1-3 m/s for liquid, 10-30m/s for gas
  - The diameter and fluid velocity should be calculated before the head loss calculation because the value of u1 and u2 is needed.
  - Pay attention to the pressure difference, especially for the plus or minus sign.
  - Consider all energy losses in the system using Bernoulli's equation.
  - Consider the flow type: turbulent or laminar, before calculating the friction factor. Use Blasius equation for turbulent flow and Poiseuille equation for laminar flow.
  - Use the following local resistance coefficient: 90 degree elbow: 0.75; 45 degree elbow: 0.35; 180 degree elbow: 1.5; tee: 1; search RAG for other local resistance coefficient.

Here are some general instruction for steam consumption calculation. If updated information is not given, follow the default settings.
  - Default latent heat is 2200kJ/kg. 

Here are some general instructions for carbon emissions calculation. If updated information is not given, follow the default settings.
  - The steam consumption is related to the heat duty of the reboiler.

Some specific requirements for this pump work calculation:      
  - Pressure at the cooling water pump inlet: 1 atm
  - Pressure at the condenser inlet pipe: 10 kPa, because the cooling water falls down and the flow is not continuous any more (Industrial experience). 
  - Straight pipe resistance (40m length)
  - 5 90 degree elbows
  - (Search RAG for a typical pump efficiency) Use 0.6 as the default efficiency.
  - The cooling water tower is set on the ground

Calculate carbon emissions for optimal conditions and provide the complete detailed calculation process in the finished chain of thought, including formula, calculation and result:
Any step that requires calculation should be done using the Python_REPL tool, including the unit conversion.
If any data is not available, please search online for the data and cite the source, especially for the energy consumption of the cooling tower's fan.

1. Calculate the head loss (unit:m) for the cooling water pump using Bernoulli's equation:stragith pipe : We/g=(p2-p1)/density/g+(u2^2-u1^2)/2g+h_f+Z2-Z1, where h_f is the friction loss
  - Calculate the energy consumption of the cooling tower's fan
  - Calculate the total energy consumption of the cooling tower
  - Pay attention to the value of u1 and u2, and actually u1=u2
2. The energy consumption of the cooling tower's fan is decided by the cooling water flow rate. Search following table for the energy consumption of the cooling tower's fan:{gxt_table_data}
3. Compute emissions using:
steam factors = {steam_factors}
4. Power emission factors:
power factors = {power_factors}
5. Calculate the carbon emission
6. For each data point, provide:
   - Source of the value (if from input data)
   - Complete calculation process (if calculated)
   - All intermediate steps and formulas used
   - Units for all values and pay attention to the unit conversion, especially for the unit conversion for carbon emission factors (kg CO2/MWh) and duty
   - 1 TJ=10^12 J 
7. Do NOT use the same power emission factor for all fuel types. Each fuel type must use its corresponding power emission factor.
8. Output format:
The complete detailed calculation process should be shown in the finished chain of thought, including formula, calculation and result
| Fuel Type | Steam Emissions (kg/h) | Power Emissions (kg/h) | Total Emissions (kg/h) | Error Range |
|---------|-----------------|-----------------|---------------|---------------|
| Coal    |-----------------|-----------------|-----------------|---------------|
| Natural Gas |-----------------|-----------------|-----------------|---------------|
| Biomass |-----------------|-----------------|-----------------|---------------|


Raw data:
{rep_content}
Input data:
{inp_content}

Note: {FORMAT_INSTRUCTIONS}"""

        print("\n\033[1;34m=== Executing Carbon Emissions Calculation ===\033[0m")
        
        try:
            response = agent.invoke({"input": prompt})
            raw_output = response['output']
            
            # Parse output
            parsed_output = raw_output
            
            print(f"\n\033[1;32mCalculation Result:\n{parsed_output}\033[0m")
            
            # Save to Word document
            save_to_word(parsed_output)
            
        except Exception as e:
            print(f"\033[1;31mCalculation failed: {str(e)}\033[0m")

    except Exception as e:
        print(f"\033[1;31mMain function execution failed: {str(e)}\033[0m")

if __name__ == "__main__":
    asyncio.run(main())