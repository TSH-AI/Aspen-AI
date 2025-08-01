import os
import re
import webbrowser
import requests
import json
from datetime import datetime
from docx import Document

# Carbon Emissions Comparison Chart Generator - Automatically extracts data from docx files
# This script reads carbon emission analysis reports for traditional distillation and heat pump assisted distillation, generates comparison charts

def extract_data_from_docx_files():
    """Extract carbon emission data from docx files"""
    
    # Get script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # File paths
    traditional_file = "inp_generate/simple/carbon_accounting/traditional_carbon_emissions.docx"
    heatpump_file = "inp_generate/difficult/carbon_accounting/heatpump_carbon_emissions.docx"
    
    # Initialize data structure
    traditional_data = {
        'steam_emissions': [0, 0, 0],  # Coal, Natural Gas, Biomass
        'power_emissions': [0, 0, 0],  # Coal, Natural Gas, Wind
        'total_emissions': [0, 0, 0],
        'error_ranges': [[0, 0], [0, 0], [0, 0]]  # [lower, upper] for each fuel type
    }
    
    heatpump_data = {
        'power_emissions': [0, 0, 0],  # Coal, Natural Gas, Wind
        'total_emissions': [0, 0, 0],
        'error_ranges': [[0, 0], [0, 0], [0, 0]]  # [lower, upper] for each fuel type
    }
    
    # Extract traditional distillation data
    try:
        doc = Document(traditional_file)
        traditional_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
        # Find emission results table
        table_pattern = r'\| Fuel Type\s+\| Steam Emissions \(kg/h\)\s+\| Power Emissions \(kg/h\)\s+\| Total Emissions \(kg/h\)\s+\| Error Range\s+\|(.*?)(?=\n\n|\Z)'
        table_match = re.search(table_pattern, traditional_content, re.DOTALL)
        
        if table_match:
            table_content = table_match.group(1)
            
            # Parse each line of data
            lines = table_content.strip().split('\n')
            for line in lines:
                if '|' in line and not line.strip().startswith('|---'):
                    parts = [part.strip() for part in line.split('|') if part.strip()]
                    if len(parts) >= 5:
                        fuel_type = parts[0].lower()
                        steam_emission = float(parts[1])
                        power_emission = float(parts[2])
                        total_emission = float(parts[3])
                        error_range = parts[4]
                        
                        # Parse error range
                        error_match = re.search(r'(\d+\.?\d*)\s*-\s*(\d+\.?\d*)', error_range)
                        if error_match:
                            error_lower = float(error_match.group(1))
                            error_upper = float(error_match.group(2))
                        else:
                            error_lower = error_upper = total_emission
                        
                        # Assign to corresponding fuel type
                        if 'coal' in fuel_type:
                            traditional_data['steam_emissions'][0] = steam_emission
                            traditional_data['power_emissions'][0] = power_emission
                            traditional_data['total_emissions'][0] = total_emission
                            # Store the actual error range from the table
                            traditional_data['error_ranges'][0] = [error_lower, error_upper]
                        elif 'natural gas' in fuel_type:
                            traditional_data['steam_emissions'][1] = steam_emission
                            traditional_data['power_emissions'][1] = power_emission
                            traditional_data['total_emissions'][1] = total_emission
                            # Store the actual error range from the table
                            traditional_data['error_ranges'][1] = [error_lower, error_upper]
                        elif 'renewable' in fuel_type or 'wind' in fuel_type or 'biomass' in fuel_type:
                            traditional_data['steam_emissions'][2] = steam_emission
                            traditional_data['power_emissions'][2] = power_emission
                            traditional_data['total_emissions'][2] = total_emission
                            # Store the actual error range from the table
                            traditional_data['error_ranges'][2] = [error_lower, error_upper]
            
            print("Successfully extracted traditional distillation data:")
            print(f"  Steam emissions: {traditional_data['steam_emissions']}")
            print(f"  Power emissions: {traditional_data['power_emissions']}")
            print(f"  Total emissions: {traditional_data['total_emissions']}")
            print(f"  Error ranges: {traditional_data['error_ranges']}")
            
    except FileNotFoundError:
        print(f"Error: Cannot find file {traditional_file}")
        return None, None
    except Exception as e:
        print(f"Error: Error reading traditional distillation file: {e}")
        return None, None
    
    # Extract heat pump assisted distillation data
    try:
        doc = Document(heatpump_file)
        heatpump_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
        # Find emission results table
        table_pattern = r'\| Fuel Type\s+\| Power Emissions \(kg/h\)\s+\| Error Range\s+\|(.*?)(?=\n\n|\Z)'
        table_match = re.search(table_pattern, heatpump_content, re.DOTALL)
        
        if table_match:
            table_content = table_match.group(1)
            
            # Parse each line of data
            lines = table_content.strip().split('\n')
            for line in lines:
                if '|' in line and not line.strip().startswith('|---'):
                    parts = [part.strip() for part in line.split('|') if part.strip()]
                    if len(parts) >= 3:
                        fuel_type = parts[0].lower()
                        power_emission = float(parts[1])
                        error_range = parts[2]
                        
                        # Parse error range
                        error_match = re.search(r'(\d+\.?\d*)\s*-\s*(\d+\.?\d*)', error_range)
                        if error_match:
                            error_lower = float(error_match.group(1))
                            error_upper = float(error_match.group(2))
                        else:
                            error_lower = error_upper = power_emission
                        
                        # Assign to corresponding fuel type
                        if 'coal' in fuel_type:
                            heatpump_data['power_emissions'][0] = power_emission
                            heatpump_data['total_emissions'][0] = power_emission
                            # Store the actual error range from the table
                            heatpump_data['error_ranges'][0] = [error_lower, error_upper]
                        elif 'natural gas' in fuel_type:
                            heatpump_data['power_emissions'][1] = power_emission
                            heatpump_data['total_emissions'][1] = power_emission
                            # Store the actual error range from the table
                            heatpump_data['error_ranges'][1] = [error_lower, error_upper]
                        elif 'wind' in fuel_type:
                            heatpump_data['power_emissions'][2] = power_emission
                            heatpump_data['total_emissions'][2] = power_emission
                            # Store the actual error range from the table
                            heatpump_data['error_ranges'][2] = [error_lower, error_upper]
            
            print("Successfully extracted heat pump assisted distillation data:")
            print(f"  Power emissions: {heatpump_data['power_emissions']}")
            print(f"  Total emissions: {heatpump_data['total_emissions']}")
            print(f"  Error ranges: {heatpump_data['error_ranges']}")
            
    except FileNotFoundError:
        print(f"Error: Cannot find file {heatpump_file}")
        return None, None
    except Exception as e:
        print(f"Error: Error reading heat pump file: {e}")
        return None, None
    
    return traditional_data, heatpump_data

def generate_html_with_ai(traditional_data, heatpump_data):
    """Use large language model to intelligently generate HTML with Plotly charts and export functionality"""
    
    # API configuration
    API_KEY = "<your api key>"
    API_URL = "<your base url>/chat/completions"
    
    # Build prompt for complete HTML generation with export functionality
    prompt = f"""
Generate a complete HTML file for carbon emissions comparison chart using Plotly.js with stacked bar charts, error bars, and PNG/SVG export functionality.

## Carbon Emission Data:

### Traditional distillation data:
- Steam emissions: {traditional_data['steam_emissions']}
- Power emissions: {traditional_data['power_emissions']}
- Total emissions: {traditional_data['total_emissions']}
- Error ranges (lower, upper): {traditional_data['error_ranges']}

### Heat pump assisted distillation data:
- Power emissions: {heatpump_data['power_emissions']}
- Total emissions: {heatpump_data['total_emissions']}
- Error ranges (lower, upper): {heatpump_data['error_ranges']}

## Chart Requirements:

### Structure:
- Create a stacked bar chart comparing traditional distillation and heat pump-assisted distillation
- 6 bars total: 3 for traditional distillation + 3 for heat pump distillation
- X-axis: Fuel types (Coal, Natural gas, Renewables) for both processes
- Use a vertical dashed line to separate the two processes
- Add section labels above each process

### Bar Chart Design:
- Traditional distillation: Stack steam emissions (top) and power emissions (bottom)
  - Power emissions should be plotted FIRST (bottom layer) so they appear at the BOTTOM of the stacked bar
  - Steam emissions should be plotted SECOND (top layer) so they appear at the TOP of the stacked bar
  - This ensures steam emissions are visually on top and their error ranges are more prominent
- Heat pump distillation: Only power emissions (steam emissions are zero)
- Use barmode: 'stack' for proper stacking
- Bar width: 0.7 for better spacing
- Legend: Only show two entries - "Steam" and "Power" (not "Heat pump power emissions")

### Colors:
- Steam emissions: #8FAADC (darker blue)
- Power emissions: #DEEBF7 (light blue)
- Error bars: #2f2f2f (dark gray)
- Separator line: #e5e5e5 (light gray)

### Error Bars:
- Add error bars for total emissions of each process
- Traditional distillation: Use symmetric error bars with array = (upper_bound - lower_bound) / 2
- Heat pump distillation: Use asymmetric error bars with array = upper_bound - total, arrayminus = total - lower_bound
- Error bar styling: color: '#2f2f2f', thickness: 1.5, width: 8

### Layout:
- Chart dimensions: 1200x600px
- Calculate X-axis range based on data: use positions [0, 1, 2, 3, 4, 5] with appropriate padding
- Calculate Y-axis range based on data: find min/max of all emission values and add 10% padding
- X-axis labels: ['Coal', 'Natural gas', 'Renewables', 'Coal', 'Natural gas', 'Renewables']
- Y-axis title: "Carbon emissions (kg CO₂/h)"
- Font: Arial throughout
- Axis lines: #2f2f2f, width: 2
- **CRITICAL: Disable ALL grid lines** - Set showgrid: false for both xaxis and yaxis to remove all background grid lines
- **CRITICAL: Keep coordinate axis lines visible** - Set showline: true for both xaxis and yaxis to keep the coordinate axis lines
- Keep the vertical dashed separator line at x=2.5 (this is NOT a grid line, it's a custom shape)
- **CRITICAL: Set plot background to white** - Set plot_bgcolor: 'white' and paper_bgcolor: 'white' in layout configuration

### Grid Lines Configuration:
- **MUST set showgrid: false for both xaxis and yaxis** in the layout configuration
- **CRITICAL: Keep axis lines visible** - Set showline: true for both xaxis and yaxis to keep the coordinate axis lines
- This removes all background grid lines while keeping the coordinate axis lines and the custom vertical separator line
- Example layout configuration:
```javascript
layout: {{
    xaxis: {{
        showgrid: false,  // Disable x-axis grid lines
        showline: true,   // Keep x-axis line visible
        linecolor: '#2f2f2f',
        linewidth: 2,
        // ... other xaxis properties
    }},
    yaxis: {{
        showgrid: false,  // Disable y-axis grid lines
        showline: true,   // Keep y-axis line visible
        linecolor: '#2f2f2f',
        linewidth: 2,
        // ... other yaxis properties
    }}
}}
```

### Styling:
- Background: white
- Plot background: white
- Margins: left=120, right=30, top=70, bottom=120
- Font sizes: axis titles 32px, tick labels 30px, legend 28px
- Legend position: top-right (x=0.99, y=0.85)
- Disable display mode bar

### Separator and Labels:
- Vertical dashed line at x=2.5, color: #e5e5e5, width: 2
- Section labels: "Traditional distillation" and "Heat pump-assisted distillation"
- Calculate label positions based on Y-axis range: position labels near the top of the chart area
- Label font: Arial, size 32, color: #2f2f2f

### Data Arrays:
```javascript
const traditionalSteamEmissions = {traditional_data['steam_emissions']};
const traditionalPowerEmissions = {traditional_data['power_emissions']};
const traditionalTotalEmissions = {traditional_data['total_emissions']};
const traditionalErrorRanges = {traditional_data['error_ranges']};

const heatpumpPowerEmissions = {heatpump_data['power_emissions']};
const heatpumpTotalEmissions = {heatpump_data['total_emissions']};
const heatpumpErrorRanges = {heatpump_data['error_ranges']};
```

### CRITICAL Stacking Order Instructions:
- For traditional distillation: Add power emissions trace FIRST (bottom layer), then add steam emissions trace SECOND (top layer)
- This ensures steam emissions appear visually on top of the stacked bar
- The steam emissions are the primary source of uncertainty, so they should be prominently displayed at the top
- Legend: Only show two entries - "Steam" and "Power". Use showlegend: false for heat pump traces to avoid duplicate legend entries

### EXPORT FUNCTIONALITY:
Add export buttons and functionality for PNG and SVG formats:

1. Add export buttons above the chart:
```html
<div style="text-align: center; margin: 20px 0;">
    <button onclick="exportChart('png')" style="background: #4CAF50; color: white; border: none; padding: 10px 20px; margin: 0 10px; border-radius: 5px; cursor: pointer; font-size: 16px;">Export as PNG</button>
    <button onclick="exportChart('svg')" style="background: #2196F3; color: white; border: none; padding: 10px 20px; margin: 0 10px; border-radius: 5px; cursor: pointer; font-size: 16px;">Export as SVG</button>
</div>
```

2. Add JavaScript export function:
```javascript
function exportChart(format) {{
    const chart = document.getElementById('chart');
    const timestamp = new Date().toISOString().replace(/[:.]/g, '-').slice(0, 19);
    const filename = `carbon_emissions_comparison_${{timestamp}}.${{format}}`;
    
    Plotly.downloadImage(chart, {{
        format: format,
        filename: filename,
        height: 600,
        width: 1200,
        scale: 2  // High resolution
    }});
}}
```

3. Ensure the chart div has id="chart"

### Technical Requirements:
- Use Plotly.js version 2.24.1
- Include complete DOCTYPE, head, body sections
- Responsive design with max-width 1200px
- Professional scientific presentation style
- Add export functionality with high-resolution output

Please generate a complete HTML file that creates a professional stacked bar chart with error bars comparing the two distillation processes, including PNG and SVG export functionality.
"""

    try:
        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "<your model>",
            "messages": [
                {
                    "role": "system",
                    "content": "You are an expert HTML/JavaScript generator for scientific data visualization. Create professional Plotly.js charts with precise data handling and export functionality. Generate complete, production-ready HTML files with stacked bar charts, error bars, and PNG/SVG export buttons. Use exact colors, fonts, and styling as specified. CRITICAL: Calculate axis ranges dynamically based on data min/max values with appropriate padding. CRITICAL: Create high-quality visualizations with proper spacing, clear markers, and professional appearance. CRITICAL: For stacked bar charts, the FIRST trace added appears at the BOTTOM and the LAST trace added appears at the TOP. To show steam emissions on top, add the power emissions trace FIRST (bottom layer), then add the steam emissions trace SECOND (top layer). CRITICAL: Legend should only show two entries - 'Steam' and 'Electricity'. Use showlegend: false for heat pump traces to avoid duplicate legend entries. CRITICAL: Include export functionality with PNG and SVG download buttons that generate high-resolution images. CRITICAL: Disable ALL grid lines by setting showgrid: false for both xaxis and yaxis in the layout configuration. CRITICAL: Keep coordinate axis lines visible by setting showline: true for both xaxis and yaxis. The vertical dashed separator line at x=2.5 should be kept as a custom shape, not as a grid line. CRITICAL: Set plot background to white by setting plot_bgcolor: 'white' and paper_bgcolor: 'white' in the layout configuration."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ],
            "max_tokens": 8000,
            "temperature": 0.1
        }
        
        print("Using AI to generate chart with export functionality...")
        response = requests.post(API_URL, headers=headers, json=data, timeout=60)
        
        if response.status_code == 200:
            result = response.json()
            html_content = result['choices'][0]['message']['content']
            
            # Clean response and extract HTML content
            if '```html' in html_content:
                html_content = html_content.split('```html')[1].split('```')[0].strip()
            elif '```' in html_content:
                html_content = html_content.split('```')[1].split('```')[0].strip()
            
            print("AI successfully generated chart HTML code with export functionality!")
            return html_content
            
        else:
            print(f"AI API call failed: {response.status_code}")
            print(f"Error message: {response.text}")
            return None
            
    except Exception as e:
        print(f"Error during AI generation: {e}")
        return None

def save_and_open_html(html_content):
    """Save HTML content to file and open in browser"""
    # Use the specified output directory
    output_dir = "inp_generate/difficult"
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename without timestamp
    filename = "carbon_emissions_comparison.html"
    filepath = os.path.join(output_dir, filename)
    
    # Write content to file
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print(f"HTML file saved to: {filepath}")
    
    # Open in default browser
    webbrowser.open(f'file://{filepath}')
    
    return filepath

def main():
    print("Extracting carbon emission data from docx files...")
    traditional_data, heatpump_data = extract_data_from_docx_files()
    
    if not traditional_data or not heatpump_data:
        print("Error: Unable to extract valid data")
        return
    
    print("Using AI to generate HTML visualization with export functionality...")
    html_content = generate_html_with_ai(traditional_data, heatpump_data)
    
    if not html_content:
        print("Error: AI generation failed")
        return
    
    print("Saving and opening HTML file...")
    save_and_open_html(html_content)
    print("Complete! Chart now includes PNG and SVG export functionality.")

if __name__ == "__main__":
    main() 